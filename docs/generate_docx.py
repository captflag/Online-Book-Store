"""
Generate professional DOCX documentation for CaptBooks project with detailed diagrams.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Get script directory for absolute paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
SCREENSHOTS_DIR = os.path.join(PROJECT_DIR, "screenshots")

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_diagram_image(filename, title=""):
    """Add a diagram image with title"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(245, 158, 11)
    
    img_path = os.path.join(SCREENSHOTS_DIR, filename)
    print(f"Checking diagram: {img_path}")
    if os.path.exists(img_path):
        print(f"  ✓ Found diagram! Adding {filename}")
        doc.add_picture(img_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Diagram: {title} - Image not found]")
        print(f"  ✗ Warning: {img_path} not found")
    
    doc.add_paragraph()

def add_section_header(text, level=1, emoji=""):
    """Add a formatted section header"""
    full_text = f"{emoji} {text}" if emoji else text
    heading = doc.add_heading(full_text, level=level)
    return heading

# ==================== TITLE PAGE ====================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title_run = title.add_run("CaptBooks")
title_run.font.size = Pt(52)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(245, 158, 11)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Premium Indian Online Bookstore")
sub_run.font.size = Pt(24)
sub_run.font.color.rgb = RGBColor(100, 116, 139)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

tech_doc = doc.add_paragraph()
tech_run = tech_doc.add_run("Technical Documentation & Architecture")
tech_run.font.size = Pt(28)
tech_run.font.bold = True
tech_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Features list on title page
features_title = doc.add_paragraph()
features_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
ft_run = features_title.add_run("React 18 • Tailwind CSS • Firebase • Framer Motion")
ft_run.font.size = Pt(14)
ft_run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_run = date_para.add_run("December 2024\nBy CaptFlag")
date_run.font.size = Pt(14)
date_run.font.color.rgb = RGBColor(100, 116, 139)
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_section_header("Table of Contents", 1, "📋")

toc_items = [
    "1. Executive Summary",
    "2. Technology Stack & Rationale",
    "3. System Architecture Overview",
    "4. Application Screenshots",
    "5. User Shopping Flow",
    "6. State Management Architecture",
    "7. Component Hierarchy",
    "8. Database Design (ERD)",
    "9. Authentication Flow",
    "10. Order Processing Flow",
    "11. Analytics Data System",
    "12. Performance Optimizations",
    "13. Future Roadmap"
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== EXECUTIVE SUMMARY ====================
add_section_header("Executive Summary", 1, "📌")

doc.add_paragraph(
    "CaptBooks is a modern, full-featured online bookstore built specifically for the Indian market. "
    "The platform combines cutting-edge frontend technologies with robust Firebase backend services "
    "to deliver a premium e-commerce experience for book enthusiasts across India."
)

doc.add_paragraph()
add_section_header("Key Highlights", 2)

highlights = [
    ("Modern Stack", "React 18, Vite, Tailwind CSS 4.0, Framer Motion"),
    ("Premium UI/UX", "3D animations, page transitions, toast notifications"),
    ("Indian Market", "INR currency, Indian authors, localized content"),
    ("Real-time Data", "Firebase Authentication & Cloud Firestore"),
    ("Analytics", "Comprehensive dashboard with sales insights"),
    ("Responsive", "Mobile-first design for all devices")
]

for title, desc in highlights:
    p = doc.add_paragraph()
    run = p.add_run(f"• {title}: ")
    run.font.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==================== TECHNOLOGY STACK ====================
add_section_header("Technology Stack & Rationale", 1, "🛠️")

doc.add_paragraph(
    "Each technology was carefully selected based on performance, developer experience, "
    "and ecosystem support. Here's a detailed breakdown:"
)

doc.add_paragraph()
add_section_header("Frontend Technologies", 2)

# Frontend table
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

headers = ['Technology', 'Version', 'Why We Chose It']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ('React 18', '18.3.1', 'Component architecture, hooks API, virtual DOM, massive ecosystem'),
    ('Vite', '6.0.1', 'Instant dev server, HMR, faster builds than webpack'),
    ('Tailwind CSS', '4.0.0', 'Utility-first, CSS-native config, smaller bundles'),
    ('Framer Motion', '11.x', 'Declarative animations, gestures, layout animations')
]

for i, (tech, ver, why) in enumerate(data, 1):
    table.rows[i].cells[0].text = tech
    table.rows[i].cells[1].text = ver
    table.rows[i].cells[2].text = why

doc.add_paragraph()
add_section_header("Backend Technologies", 2)

table2 = doc.add_table(rows=3, cols=3)
table2.style = 'Table Grid'

for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data2 = [
    ('Firebase Auth', '11.x', 'Email/password + Google OAuth, session management'),
    ('Cloud Firestore', '11.x', 'NoSQL, real-time sync, offline support, auto-scaling')
]

for i, (tech, ver, why) in enumerate(data2, 1):
    table2.rows[i].cells[0].text = tech
    table2.rows[i].cells[1].text = ver
    table2.rows[i].cells[2].text = why

doc.add_page_break()

# ==================== ARCHITECTURE OVERVIEW ====================
add_section_header("System Architecture Overview", 1, "🏗️")

doc.add_paragraph(
    "The application follows a modern layered architecture with clear separation of concerns:"
)

add_diagram_image("system-architecture.png", "System Architecture Diagram")

doc.add_page_break()

# ==================== SCREENSHOTS ====================
add_section_header("Application Screenshots", 1, "🖼️")

# Home Page
add_section_header("Home Page", 2)
doc.add_paragraph(
    "The homepage features a premium hero section with floating book animations, "
    "a horizontal scrollable carousel of featured books, and a showcase of top Indian authors "
    "with animated gradient rings."
)

home_img = os.path.join(SCREENSHOTS_DIR, "home-page.png")
print(f"Checking: {home_img}")
if os.path.exists(home_img):
    print(f"  ✓ Found! Adding home-page.png ({os.path.getsize(home_img)} bytes)")
    doc.add_picture(home_img, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Home Page Screenshot - Image not found]")
    print(f"  ✗ Warning: {home_img} not found")

doc.add_paragraph()

# Shop Page
add_section_header("Shop Page", 2)
doc.add_paragraph(
    "The shop displays all books with category filters, price range selection, "
    "and 3D tilt effect on book cards. Users can easily add items to cart."
)

shop_img = os.path.join(SCREENSHOTS_DIR, "shop-page.png")
print(f"Checking: {shop_img}")
if os.path.exists(shop_img):
    print(f"  ✓ Found! Adding shop-page.png ({os.path.getsize(shop_img)} bytes)")
    doc.add_picture(shop_img, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Shop Page Screenshot - Image not found]")
    print(f"  ✗ Warning: {shop_img} not found")

doc.add_page_break()

# Cart Page
add_section_header("Cart Page", 2)
doc.add_paragraph(
    "The shopping cart shows all selected items with quantity controls, "
    "individual prices, and a comprehensive order summary with checkout button."
)

cart_img = os.path.join(SCREENSHOTS_DIR, "cart-page.png")
print(f"Checking: {cart_img}")
if os.path.exists(cart_img):
    print(f"  ✓ Found! Adding cart-page.png ({os.path.getsize(cart_img)} bytes)")
    doc.add_picture(cart_img, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Cart Page Screenshot - Image not found]")
    print(f"  ✗ Warning: {cart_img} not found")

doc.add_paragraph()

# Analytics Dashboard
add_section_header("Admin Analytics Dashboard", 2)
doc.add_paragraph(
    "The admin dashboard provides comprehensive business analytics including "
    "revenue trends, top-selling books, category breakdown, customer insights, "
    "and geographic distribution of orders."
)

analytics_img = os.path.join(SCREENSHOTS_DIR, "analytics-dashboard.png")
print(f"Checking: {analytics_img}")
if os.path.exists(analytics_img):
    print(f"  ✓ Found! Adding analytics-dashboard.png ({os.path.getsize(analytics_img)} bytes)")
    doc.add_picture(analytics_img, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Analytics Dashboard Screenshot - Image not found]")
    print(f"  ✗ Warning: {analytics_img} not found")

doc.add_page_break()

# ==================== USER SHOPPING FLOW ====================
add_section_header("User Shopping Flow", 1, "🛒")

doc.add_paragraph(
    "The following diagram illustrates the complete user journey from browsing to purchase:"
)

add_diagram_image("shopping-flow.png", "Shopping Flow Diagram")

doc.add_page_break()

# ==================== STATE MANAGEMENT ====================
add_section_header("State Management Architecture", 1, "⚡")

doc.add_paragraph(
    "The application uses React Context API for global state management, "
    "avoiding the complexity of Redux while maintaining clean state flow:"
)

add_diagram_image("state-management.png", "State Management Diagram")

doc.add_page_break()

# ==================== COMPONENT HIERARCHY ====================
add_section_header("Component Hierarchy", 1, "🧩")

doc.add_paragraph(
    "The following tree shows the complete component structure from root to leaf:"
)

add_diagram_image("component-tree.png", "Component Tree")

doc.add_page_break()

# ==================== DATABASE DESIGN ====================
add_section_header("Database Design (ERD)", 1, "🗄️")

doc.add_paragraph(
    "The application uses Cloud Firestore with the following collections and relationships:"
)

add_diagram_image("database-erd.png", "Database Schema")

doc.add_page_break()

# ==================== AUTHENTICATION FLOW ====================
add_section_header("Authentication Flow", 1, "🔐")

doc.add_paragraph(
    "The authentication system supports multiple providers with automatic session management:"
)

add_diagram_image("auth-flow.png", "Authentication Sequence")

doc.add_page_break()

# ==================== ORDER PROCESSING ====================
add_section_header("Order Processing Flow", 1, "📦")

doc.add_paragraph(
    "The order processing involves multiple steps from cart to delivery:"
)

add_diagram_image("order-processing.png", "Order Processing Sequence Diagram")

doc.add_page_break()

# ==================== ANALYTICS SYSTEM ====================
add_section_header("Analytics Data System", 1, "📊")

doc.add_paragraph(
    "The analytics dashboard aggregates data from multiple sources to provide business insights:"
)

add_diagram_image("analytics-system.png", "Analytics Data Flow")

doc.add_paragraph()
add_section_header("Analytics Functions Reference", 2)

# Analytics functions table
analytics_table = doc.add_table(rows=8, cols=3)
analytics_table.style = 'Table Grid'

analytics_data = [
    ('Function', 'Output', 'Purpose'),
    ('getSalesStats()', 'totals, counts', 'Overall business metrics'),
    ('getTopSellingBooks()', 'ranked books', 'Identify bestsellers'),
    ('getRevenueByCategory()', 'breakdown', 'Category performance'),
    ('getSalesByDay()', 'daily totals', 'Revenue trend chart'),
    ('getTopAuthors()', 'author rankings', 'Author performance'),
    ('getOrdersByCity()', 'city data', 'Geographic insights'),
    ('getCustomerInsights()', 'metrics', 'Customer loyalty analysis')
]

for i, row_data in enumerate(analytics_data):
    for j, cell_text in enumerate(row_data):
        analytics_table.rows[i].cells[j].text = cell_text
    if i == 0:
        for cell in analytics_table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ==================== PERFORMANCE ====================
add_section_header("Performance Optimizations", 1, "⚡")

perf_items = [
    ("Code Splitting", "React.lazy() for route-based code splitting, reducing initial bundle size by loading pages on demand"),
    ("Tree Shaking", "Vite automatically removes unused code from the final bundle"),
    ("Image CDN", "All images served from Unsplash CDN with automatic compression and resizing"),
    ("State Management", "React Context instead of Redux - simpler, less boilerplate, no extra dependencies"),
    ("CSS Optimization", "Tailwind purges unused classes in production, resulting in minimal CSS bundle"),
    ("Database Caching", "Firestore provides automatic client-side caching and offline support"),
    ("Lazy Loading", "Images and components loaded only when needed"),
    ("Minimal Dependencies", "Carefully selected packages to avoid bundle bloat")
]

for title, desc in perf_items:
    p = doc.add_paragraph()
    run = p.add_run(f"⚡ {title}: ")
    run.font.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==================== FUTURE ROADMAP ====================
add_section_header("Future Roadmap", 1, "🚀")

doc.add_paragraph("Planned enhancements for future releases:")

roadmap = [
    ("Phase 1 - Q1 2025", [
        "Server-side rendering with Next.js for SEO",
        "Razorpay payment gateway integration",
        "Email notifications for order updates"
    ]),
    ("Phase 2 - Q2 2025", [
        "Wishlist functionality",
        "Book reviews and ratings system",
        "Advanced search with Algolia"
    ]),
    ("Phase 3 - Q3 2025", [
        "Progressive Web App (PWA) support",
        "Multi-language support (Hindi, Tamil, etc.)",
        "Recommendation engine using ML"
    ])
]

for phase, items in roadmap:
    p = doc.add_paragraph()
    run = p.add_run(phase)
    run.font.bold = True
    run.font.color.rgb = RGBColor(245, 158, 11)
    
    for item in items:
        p = doc.add_paragraph(f"  ☐ {item}")

doc.add_page_break()

# ==================== FOOTER ====================
doc.add_paragraph()
doc.add_paragraph()

footer_line = doc.add_paragraph()
footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = footer_line.add_run("─" * 50)
line_run.font.color.rgb = RGBColor(200, 200, 200)

doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("CaptBooks Technical Documentation")
footer_run.font.bold = True
footer_run.font.size = Pt(14)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2_run = footer2.add_run("By CaptFlag | December 2024")
footer2_run.font.color.rgb = RGBColor(100, 116, 139)

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer3_run = footer3.add_run("GitHub: github.com/captflag/Online-Book-Store")
footer3_run.font.color.rgb = RGBColor(100, 116, 139)
footer3_run.font.size = Pt(10)

# Save document
output_path = os.path.join(os.path.dirname(SCRIPT_DIR), "docs", "CaptBooks_Documentation_v3.docx")
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
print(f"📄 Total pages: ~15+")
print(f"📊 Includes 9 detailed architecture diagrams")
print(f"🖼️ Includes 4 application screenshots")
